from grab import Grab
from docx import Document
import time
import urllib.request

# =================================
# =================================
# INPUT
main_html_file_path = ""
channel_name = ""

# OUTPUT
docx_filename = ""
# =================================
# =================================

if len(main_html_file_path) < 2 or len(channel_name) < 2 or not main_html_file_path.endswith(".html"):
    print("Please correctly fill 'main_html_file_path' and 'channel_name' fields before using the script!")
    exit(1)

if len(docx_filename) < 2:
    docx_filename = channel_name

if not docx_filename.endswith(".docx"):
    docx_filename += ".docx"

document = Document()
document.add_heading(channel_name, 0)
document.add_heading('', 0)

with open(main_html_file_path, 'rb') as fobj:
    f = fobj.read()

g = Grab()
g.setup_document(f)

all_publications = g.doc.select('//div[@class="feed__item-wrap _size_small"]')

xpath_name = '//div[@class="clamp__text-expand"]'
xpath_link = '//a[@class="card-image-view__clickable"]//@href'
xpath_publication_body = '//div[@class="article__middle"]'
xpath_publication_text = '//div[@class="article__body"]//p'

print("all publication nums: ", len(all_publications))

i = 0
for publication in all_publications:

    try:
        publication_name = publication.select(xpath_name)[i].text()
        publication_link = publication.select(xpath_link)[i].text()
    except Exception as e:
        print("\tError getting publication name or link: " + str(e))
        continue

    print(str(i) + ":", publication_name, "\n\tlink: ", publication_link)

    document.add_heading(publication_name, level=1)

    print("\tGetting publication " + str(i) + " by link...")

    html = urllib.request.urlopen(publication_link).read()
    g_publication = Grab()
    g_publication.setup_document(html)

    try:
        publication_body = g_publication.doc.select(xpath_publication_body)[0]
        publication_multi_text = publication_body.select(xpath_publication_text)
    except Exception as e:
        print("\tError: " + str(e))
        print("\tThis link maybe not be publication!")
        i += 1
        continue

    print("\tNumber of paragraph in publication" + str(i) + ": " + str(len(publication_multi_text)-1))

    publication_to_file = ""

    for one_paragraph in publication_multi_text:
        paragraph_in_str = one_paragraph.text()

        if len(paragraph_in_str) < 2:
            continue

        publication_to_file += "\t"
        publication_to_file += paragraph_in_str
        publication_to_file += "\n"

    print("\tSaving publication " + str(i) + " on disk... ")
    document.add_paragraph(publication_to_file)
    document.add_heading('', 0)
    document.save(docx_filename)

    print("\tAnti-сaptcha waiting...")
    time.sleep(5)

    print()
    i += 1